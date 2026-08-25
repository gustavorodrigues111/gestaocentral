# /api/documento-preencher — preenche um modelo trabalhista (.docx tokenizado com
# {{CAMPOS}}) do acervo do escritório e devolve { docxBase64, faltando }. Porte
# fiel do preencher.py (mesmo python-docx). NÃO gera PDF (o PDF exato sai pela
# skill/LibreOffice). Modelos em _docrh/modelos/. Recebe POST { modeloId, dados }.
import os, io, re, json, copy, base64
from http.server import BaseHTTPRequestHandler

from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH

MODELOS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_docrh", "modelos")
TOKEN = re.compile(r'\{\{([A-Z0-9_]+)\}\}')
ASSIN = re.compile(r'ASSINATURA\s+(EMPREGADOR[AO]?|EMPREGADO)\b', re.I)
ID_OK = re.compile(r'^[a-z0-9][a-z0-9-]{0,80}$')


def iter_paragraphs(doc):
    def walk(container):
        for p in container.paragraphs:
            yield p
        for t in container.tables:
            for row in t.rows:
                for cell in row.cells:
                    yield from walk(cell)
    yield from walk(doc)


def run_map(p):
    out, pos = [], 0
    for r in p.runs:
        n = len(r.text)
        out.append((r, pos, pos + n))
        pos += n
    return out


def replace_span(p, start, end, novo):
    primeiro = True
    for r, rs, re_ in run_map(p):
        if re_ <= start or rs >= end:
            continue
        a = max(0, start - rs)
        b = min(len(r.text), end - rs)
        if primeiro:
            r.text = r.text[:a] + novo + r.text[b:]
            primeiro = False
        else:
            r.text = r.text[:a] + r.text[b:]
    return not primeiro


def _clone_abaixo(p, texto, tamanho=Pt(9.5), negrito=False):
    novo = copy.deepcopy(p._p)
    p._p.addnext(novo)
    np = Paragraph(novo, p._parent)
    for r in np.runs[1:]:
        r._r.getparent().remove(r._r)
    if not np.runs:
        np.add_run("")
    r = np.runs[0]
    r.text = texto
    r.bold = negrito
    r.font.size = tamanho
    r.underline = False
    return np


def _nomes_assinatura(doc, mapa):
    if not mapa:
        return
    for p in list(iter_paragraphs(doc)):
        m = ASSIN.search(p.text or "")
        if not m:
            continue
        alvo = m.group(1).upper()
        chave = "empregadora" if alvo.startswith("EMPREGADOR") and alvo != "EMPREGADO" else "empregado"
        valor = mapa.get(chave)
        if valor:
            _clone_abaixo(p, valor)


def _bloco_testemunhas(doc):
    ref = doc.paragraphs[-1]
    linhas = [
        ("TESTEMUNHAS (na hipótese de recusa de assinatura pelo(a) empregado(a)):", True),
        ("", False),
        ("1. ____________________________________________   Nome / CPF", False),
        ("", False),
        ("2. ____________________________________________   Nome / CPF", False),
    ]
    for texto, negrito in reversed(linhas):
        np = _clone_abaixo(ref, texto, tamanho=Pt(9.5), negrito=negrito)
        np.alignment = WD_ALIGN_PARAGRAPH.LEFT


def preencher_bytes(modelo_path, dados):
    doc = Document(modelo_path)
    faltando = set()

    for p in iter_paragraphs(doc):
        while True:
            m = TOKEN.search(p.text)
            if not m:
                break
            campo = m.group(1)
            valor = dados.get(campo)
            if valor is None or str(valor).strip() == "":
                faltando.add(campo)
                valor = "____________"
            if not replace_span(p, m.start(), m.end(), str(valor)):
                break

    # Marcações ( ) → (X): acha o parágrafo/célula pela âncora e marca o 1º ( ).
    def _norm(s):
        return re.sub(r'\s+', ' ', (s or '').replace('\xa0', ' ')).strip().lower()
    box = re.compile(r'\(\s+\)')
    for mk in (dados.get("_marcar") or []):
        anc = _norm(mk.get("ancora", ""))
        if not anc:
            continue
        for p in iter_paragraphs(doc):
            if anc in _norm(p.text):
                m = box.search(p.text)
                if m:
                    replace_span(p, m.start(), m.end(), "(X)")
                break

    for ins in (dados.get("_inserir") or []):
        alvo, texto = str(ins.get("apos", "")).strip().lower(), ins.get("texto", "")
        if not alvo or not texto:
            continue
        for p in iter_paragraphs(doc):
            if alvo in p.text.strip().lower():
                novo = copy.deepcopy(p._p)
                p._p.addnext(novo)
                np = Paragraph(novo, p._parent)
                for r in np.runs[1:]:
                    r._r.getparent().remove(r._r)
                if np.runs:
                    np.runs[0].text = texto
                    np.runs[0].bold = False
                break

    # Quadros repetíveis: escreve valores nas células (linha/coluna inicial),
    # adicionando linha quando faltar. Preserva o estilo da célula.
    def _set_cell(cell, text):
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(text)
    for tb in (dados.get("_tabela") or []):
        ti = tb.get("tabela")
        if ti is None or not isinstance(ti, int) or ti < 0 or ti >= len(doc.tables):
            continue
        t = doc.tables[ti]
        li = int(tb.get("linha_inicial", 1) or 0)
        ci = int(tb.get("col_inicial", 0) or 0)
        for i, linha in enumerate(tb.get("linhas") or []):
            r = li + i
            while r >= len(t.rows):
                t.add_row()
            row = t.rows[r]
            for j, val in enumerate(linha or []):
                c = ci + j
                if c < len(row.cells) and str(val).strip():
                    _set_cell(row.cells[c], str(val))

    _nomes_assinatura(doc, dados.get("_assinaturas") or {})
    if dados.get("_testemunhas"):
        _bloco_testemunhas(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), sorted(faltando)


class handler(BaseHTTPRequestHandler):
    def _send(self, code, obj):
        body = json.dumps(obj, ensure_ascii=False).encode("utf-8")
        self.send_response(code)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_POST(self):
        try:
            n = int(self.headers.get("Content-Length") or 0)
            raw = self.rfile.read(n) if n else b""
            body = json.loads(raw.decode("utf-8")) if raw else {}
        except Exception as e:
            return self._send(400, {"error": f"JSON inválido: {e}"})

        modelo_id = str(body.get("modeloId") or "").strip()
        dados = body.get("dados") or {}
        if not ID_OK.match(modelo_id):
            return self._send(400, {"error": "modeloId inválido."})
        if not isinstance(dados, dict):
            return self._send(400, {"error": "dados deve ser um objeto."})
        path = os.path.join(MODELOS_DIR, modelo_id + ".docx")
        if not os.path.isfile(path):
            return self._send(404, {"error": f"Modelo não encontrado: {modelo_id}"})

        try:
            docx_bytes, faltando = preencher_bytes(path, dados)
        except Exception as e:
            return self._send(500, {"error": f"Falha ao preencher: {e}"})

        return self._send(200, {
            "docxBase64": base64.b64encode(docx_bytes).decode("ascii"),
            "faltando": faltando,
        })
